import docx
import json
import re
from lxml import etree
from docx_utils.parse_paper import processPaper
from MsEquation.MsEquation2Latex import get_latex
from docx_utils.ti2html import get_ti_content

### 获取标签
def get_tag(el):
    tag = el.tag
    e_index = tag.find('}')
    tag = tag[e_index + 1:]
    tag = tag.split('\s{1}')[0]
    return tag

### 获取标签内的内容
def get_content(el):
    el_tag = get_tag(el)
    text = ''
    if el_tag == 'r':
        text_el = el.xpath('.//w:t', namespaces=el.nsmap)
        if len(text_el)>0:
            text = text_el[0].text
    elif el_tag == 'oMath':
        omath = get_latex(el)
        omath ='$$'+ omath.replace('<' , '&lt;') +'$$'
        text = '<span>' + omath + '</span>'
    elif el_tag == 'drawing':
        text = 'pic'

    return text

### 获取title
def get_title(doc , title_index):
    title = ''
    for index in title_index:
        index_tree = etree.fromstring(doc.paragraphs[index]._element.xml)
        children = index_tree.getchildren()
        for child in children:
            res = get_content(child)
            title += res

    return title

### 获取options
def get_options(ops):
    options = []

    for item in ops:
        temp_ops = dict()
        res = re.findall(r'([A-Z])[.．]\s{0,1}(.*)' , item)
        temp_ops['label'] = res[0][0]
        temp_ops['content'] = res[0][1]
        options.append(temp_ops)

    return options



### 生成question
def create_question(type , stem , solution , options):
    temp_que = dict()
    temp_que["type"] = type
    temp_que["stem"] = '<span>'+ stem +'</span>'
    temp_que["solution"] = solution
    temp_que["options"] = options

    return temp_que

### 生成subject
def create_subject(title , category , reference , questions):
    temp_sub = dict()
    temp_sub["title"] = title
    temp_sub["category"] = category
    temp_sub["reference"] = reference
    temp_sub["questions"] = questions

    return temp_sub

### 选择题处理
def handle_choice(choice_sub):
    title = ''
    reference = ''
    questions = []
    category = '单项选择'

    solution = ''
    type = 'SINGLE'
    stem = choice_sub['title']
    options = get_options(choice_sub['options'])
    temp_question = create_question(type , stem , solution , options)
    questions.append(temp_question)

    subject = create_subject(title , category , reference , questions)

    return subject

### 非选择题处理
def handle_non_choice(sub):
    subject = dict()
    tit_list = sub['title'].split('<br/>')

    title = ''
    reference = ''
    solution = ''
    options = []
    type = 'GENERAL'
    stem_list = []
    questions = []

    if len(tit_list) == 1:
        stem_list.append(tit_list[0])
    else:
        i = 0
        flag = 0
        text = ''
        while 1:
            if re.match('^(\(|\（)\d{1,2}(\)|\）)', tit_list[i]) is None: # 不是以序号开头
                pass
            else:   # 以序号开头
                if flag == 0:
                    flag = 1
                    title = text
                    text = ''
                else:
                    stem_list.append(text)
                    text = ''

            text += tit_list[i]
            i+=1

            if(i==len(tit_list)):
                if flag == 0:
                    title = ''
                stem_list.append(text)
                break

    category = '填空题' if len(title)==0 else '解答题'

    for stem in stem_list:
        temp_que = create_question(type , stem , solution , options)
        questions.append(temp_que)

    subject = create_subject(title ,category , reference , questions )

    return subject

def handle_sub(sub):
    if 'options' not in sub:    #没有options选项
        sub = handle_non_choice(sub)
    else:                   # 有options选项
        sub = handle_choice(sub)

    return sub

def get_answer(doc , question_indexes):
    all_ans = []
    for que_index in question_indexes:
        ans_tit = que_index[0]
        ans_indexes = que_index[1]
        curr_index = 0
        while curr_index<len(ans_indexes):
            curr_index , ans = get_ti_content(doc , ans_indexes , curr_index ,ans_tit, '')
            content = re.sub(r'\d{1,2}[.．]\s{0,}', '', ans['title'])
            all_ans.append(content)

    return all_ans

def merge_answer(subjects , ans_list):
    for i in range(0 , len(subjects)-1):
        sub = subjects[i]
        sub['reference'] = ans_list[i].strip()
        if len(sub['questions']) == 1:
            sub['questions'][0]['solution'] = ans_list[i].strip()

    return subjects


if __name__ == "__main__":

    path = 'src/2019年全国I卷理科数学高考真题.docx'
    ans_path = 'src/2019年全国I卷理科数学高考真题答案.docx'
    doc = docx.Document(path)
    ans_doc = docx.Document(ans_path)
    paragraphs = doc.paragraphs
    all_subject = processPaper(doc)
    all_ans = processPaper(ans_doc)

    index = 0

    i = 0
    mode_text = r'完成\d～\d题'  ##模式字符串
    subjects = []

    for subject in all_subject:  ##每个大题
        sub_tit = subject[0]    # 题型段落index
        sub_list = subject[1]   # 题目集合
        curr_sub_index = 0      # 当前题目的index
        curr_question = sub_list[0]  #  当前题目
        curr_index =0

        while curr_sub_index<len(sub_list):
            curr_sub_index , sub = get_ti_content(doc , sub_list , curr_sub_index , sub_tit , mode_text)
            sub = handle_sub(sub)
            subjects.append(sub)


    ans_list = get_answer(ans_doc, all_ans)
    subjects = merge_answer(subjects , ans_list)

    with  open('math_data.json', 'w', encoding='utf-8') as fp:
        json.dump(subjects, fp, ensure_ascii=False,indent = 4, separators=(',', ': '))
