import docx
import re
#判断是否为标题
def isNumber(char):

    num1=['0', '1', '2', '3','4','5','6','7','8','9']
    num2=['一','二','三','四','五','六','七','八','九']
    if char in num1 or char in num2:
        return True
    return 0

##获取下一个数字
def get_next_number(text):
    num1=['0', '1', '2', '3','4','5','6','7','8','9']
    num2=['一','二','三','四','五','六','七','八','九']
    next_number=''
    if text.isdigit():
        next_number=str( int(text)+1)
    else:
        for i in range(0, len(num2)):
            if num2[i]==text:
                next_number = num2[i+1]
    return next_number
###计算模式字符串
def get_mode_string(text):
    tt=text[0:3]
    k=0
    mode_string=''
    for i in range(0,3):
        if isNumber(text[i]):
            k=i
            break
    if k==0:
        if text[1] in ['\t', ' ']:
            mode_string=''
    return mode_string

def analys_layout(doc):
    '''

    :param doc:
    :return: tree
    格式：[ (row,paragraphs[row].text, mode_text),(...)]
    '''
    pars=[]   ###可能
    b=0
    e=0
    paragraphs=doc.paragraphs
    row=0
    results=[]
    ##找出所有带数字的行和行号(  中文数字 一， 阿拉伯数字 1 )
    while(row<len(paragraphs)):
        text=paragraphs[row].text.strip()
        for n in range(0,3):
            if n>=len(text):
                continue
            if isNumber(text[n]):
                pars.append((n, row, text)) ##n--数字出现的位置，i--段落号
                break
        row=row+1

    ##找所有模式，行号为 1 或者 一
    modes=[]
    for n, row,text in pars:
        if n==0:
            if text[n] in ['1','一'] and (not isNumber(text[n+1])):
                modes.append((n,row, text, text[n]) )
        else:
            if text[n] in ['1','一'] and (not isNumber(text[n-1])) and (not isNumber(text[n+1])):
                modes.append(( n,row, text, text[n]) )

    tree=[]
    for i in range(0,len(modes)):  ##每个模式
        n, begin_row, text, mode_text=modes[i]  ##m为行号
        tmp_mode=[]

        tmp_mode.append((begin_row,text, mode_text))

        #print('first:',  paragraphs[begin_row].text)
        next_number=get_next_number(mode_text)

        for row in range(begin_row+1,len(paragraphs)):   ###查找某个模式的所有的值
            if paragraphs[row].text.startswith(text[:n]+mode_text+text[n+1]): ##发现了相同的模式，退出
                break

            if paragraphs[row].text.startswith(text[:n]+next_number+text[n+1]):
                #print('next_string:', text[:n] + next_number + text[n + 1], paragraphs[row].text)
                next_number = get_next_number(next_number)
                tmp_mode.append( (row,paragraphs[row].text, text[:n]+mode_text+text[n+1]) )
        if len(tmp_mode)>1:
            tree.append(tmp_mode.copy())

    return tree

##获取1个选项，[A-G]. 形式的
def get_option(text):
    text=text.strip()
    indexs=[]
    options=[]
    for item in re.finditer(r'[A-G][\.．]', text):
        indexs.append((item.group(),item.span()))
    print('in get_option,text=', text)
    if indexs[0][1]!=(0,2):  ###校检结果，保证
        print('获取选择题选项出错，请检查试题格式')
        print('text=' , text)

    i=0
    while(i<len(indexs)):
        b=indexs[i][1][0]
        if i==len(indexs)-1:
            options.append((option_text[0], text[b:].strip()))
            break
        e=indexs[i+1][1][0]
        option_text=indexs[i][0]
        options.append((option_text[0],text[b:e].strip()))
        b=e
        i=i+1
    return options

##获取某个题型模式
def get_ti_mode(tree, mode_text, start_position):
    i=start_position
    while(i<len(tree)):

        leaf=tree[i]
        if leaf[0][1].strip().startswith(mode_text):
            print('i=',i,'leaf=', tree[i])
            return i
        i=i+1
    return None

###处理1种题型（带题目类型的）
def parse_one_titype(curr_row, next_row,xiaoti_indexs ,paragraphs ):
    tis=[]
    i=0
    while(i<len(xiaoti_indexs)):
        # r,text,mode_text=xiaoti_indexs[i]
        if xiaoti_indexs[i][0]>curr_row:
            if i==len(xiaoti_indexs)-1:
                ti = parse_ti(xiaoti_indexs, xiaoti_indexs[i][0], next_row, paragraphs)

                tis.append(ti)
                break

            if xiaoti_indexs[i+1][0]<next_row:
                ti=parse_ti(xiaoti_indexs, xiaoti_indexs[i][0], xiaoti_indexs[i+1][0],paragraphs )
            else:
                ti=parse_ti(xiaoti_indexs, xiaoti_indexs[i][0], next_row, paragraphs)
                tis.append(ti)
                break
            tis.append(ti)
        i=i+1

    return tis
####处理1个题
def isObjective( curr_row, next_row, paragraphs):
    # print('next_row=',next_row)
    for i in range(curr_row, next_row):
        text=paragraphs[i].text.strip()
        if re.match(r'[A-G][．\.]', text):
            return (True,i)
    return (False,-1)

####解析1道题
def parse_ti(xiaoti_indexs, curr_row, next_row , paragraphs):
    # curr_row=xiaoti_indexs
    objective,index= isObjective( curr_row, next_row,paragraphs)
    ti={}
    ti['title']=[]
    if objective:
        options=[]
        for i in range(curr_row, index):
            ti['title'].append(i)
        for j in range(index,next_row):
            options.append( j )
        ti['options'] = options
    else:
        for i in range(curr_row, next_row):
            ti['title'].append(i)

    print('in ti=', ti)
    print('in ti',objective,index)
    return ti


def verify_options(options):

    pass

###计算主要模式在tree的位置
##2019.9.30,主模式确定后，副模式应该在主模式之前（行号更小）
def get_main_modes(tree):
    data=[]
    i=0
    while(i<len(tree)):
        data.append((i,len(tree[i]), tree[i][0][2], tree[i][0][0] ))   ##i为
        i=i+1
    data.sort(key=lambda x:x[1], reverse=True)
    print('data=', data)

    primary_mode_index=data[0][0]  ###最长的肯定是主模式，
    primary_mode_text = data[0][2]
    min_row=data[0][-1]
    second_mode_index = data[1][0]

    if  primary_mode_text[0]!='1':
        print('试卷格式可能有问题')
        print('模式字符串是：',primary_mode_text )

    i=1
    while(i<len(data)):
        if data[i][2][0]==primary_mode_text[0]:
            i = i + 1
            continue
        if '一' in data[i][2] and data[i][-1]<min_row :
            second_mode_index=data[i][0]
            break
        i = i + 1

    return (second_mode_index, primary_mode_index)



def processPaper(doc):
    '''
    默认，试卷，题目大题是 一、 这种形式
    :param doc:
    :return:
    '''
    paragraphs=doc.paragraphs
    tree=analys_layout(doc)
    # 获取一份试卷主要的大题和主干小题的在tree里的索引

    dati_mode_index, xiao_mode_index=get_main_modes(tree)   ##试卷的主要2层模式
    print('mode_index,', dati_mode_index, xiao_mode_index)


####获取所有大题的  小题
    j=0
    tis=[]

    print('tree=', tree)
    dati_indexs=tree[dati_mode_index]
    xiaoti_indexs=tree[xiao_mode_index]

    curr_row, text, mode_text = dati_indexs[0]
    i=0
    all_ti=[]
    while(i<len(dati_indexs)):
        if i<len(dati_indexs)-1:
            next_row, next_text,mode_text=dati_indexs[i+1]
            tis=parse_one_titype(curr_row, next_row, xiaoti_indexs, paragraphs)
            # print('tis=', tis)
        else:
            tis=parse_one_titype(curr_row, len(paragraphs), xiaoti_indexs, paragraphs)
            # print('tis=', tis)
        all_ti.append(tis)
        i=i+1
        curr_row=next_row

    return all_ti

###获取题型 ---必须以(一、)(二、)开头
def get_ti_types(tree):
    i=0
    mode_index=-1
    ti_types=[]
    while(i<len(tree)):
        for row, text, mode_text in tree[i]:
            if text.strip().startswith('一、'):    ##大题肯定以 '一、'开始
                mode_index=i
                break
        if mode_index!=-1:
            break
        i=i+1

    for i in range(0,len(tree[mode_index])) :
        row, text, mode_text =tree[mode_index][i]
        x=text.find('题')
        if x!=-1:
            ti_types.append((row,text[2:x+1]))
        else:
            ti_types.append( (row, remove_brackets(text.strip()) ) )
    return (mode_index, ti_types)

####删除括号及其里面的内容
def remove_brackets(sentence):
    result=re.findall(r'(^.*)[（\(].*[\)）](.*)', sentence)
    return ''.join(result)

'''
试卷的格式，我们认为只有2级，
1.大题（一、填空题）
2. 小题（19.）

'''

if __name__ == "__main__":
    path='src/2019年全国I卷理科数学高考真题.docx'
    doc=docx.Document(path)
    all_ti_index=processPaper(doc)
