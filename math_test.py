import docx
import re
from lxml import etree
from parse_docx_paper import processPaper
from MsEquation.MsEquation2Latex import get_latex


### 生成tree
def get_tree(index):
    return etree.fromstring(doc.paragraphs[index]._element.xml)

### 获取标签
def get_tag(el):
    tag = el.tag
    e_index = tag.find('}')
    tag = tag[e_index + 1:]
    tag = tag.split('\s{1}')[0]
    return tag

### 获取标签内的内容
def get_content(el):
    el_tag = get_tag(el)
    text = ''
    if el_tag == 'r':
        text_el = el.xpath('.//w:t', namespaces=el.nsmap)
        if len(text_el)>0:
            text = text_el[0].text
    elif el_tag == 'oMath':
        omath = get_latex(el)
        omath ='$$'+ omath.replace('<' , '&lt;') +'$$'
        text = '<span>' + omath + '</span>'
    elif el_tag == 'drawing':
        text = 'pic'

    return text


### 获取title
def get_title(doc , title_index):
    title = ''
    for index in title_index:
        index_tree = etree.fromstring(doc.paragraphs[index]._element.xml)
        children = index_tree.getchildren()
        for child in children:
            res = get_content(child)
            title += res

    return title

### 获取options
def get_options(doc , ops_index):
    option_list = []
    for index in ops_index:
        ops_tree = etree.fromstring(doc.paragraphs[index]._element.xml)
        ops_children = ops_tree.getchildren()
        content = ''
        for child in ops_children:
            res = get_content(child)
            if res == '':
                res = 'br'
            content += res

    content = content.split('br')
    for op in content:
        if op:
            temp_op = dict()
            options = re.findall(r'([A-D])[.．]\s{0,1}(.+)' , op)
            temp_op["label"] = options[0][0]
            temp_op["content"] = options[0][1]
            option_list.append(temp_op)

    return option_list

### 生成question
def create_que(type , stem , solution , options):
    print(stem , '76')
    temp_que = dict()
    temp_que["type"] = type
    temp_que["stem"] = stem
    temp_que["solution"] = solution
    temp_que["options"] = options

    print(temp_que['stem'])

    exit()
    return temp_que

### 生成subject
def create_subject(title , category , reference , questions):
    temp_sub = dict()
    temp_sub["title"] = title
    temp_sub["category"] = category
    temp_sub["reference"] = reference
    temp_sub["question"] = questions

    return temp_sub


### 选择题处理
def handle_choice(doc , choice_item):
    title = ''
    reference = ''
    questions = []
    category = '单项选择'
    tit_index = choice_item["title"]
    ops_index = choice_item['options']

    solution = ''
    type = 'SINGLE'
    stem = '<span>' + get_title(doc, tit_index) +'</span>'
    options = get_options(doc , ops_index)
    temp_que = create_que(type , stem , solution , options)
    questions.append(temp_que)

    subject = create_subject(title , category , reference , questions)

    print(subject , '120')
    exit()

    return subject

def paraghraph2html(paragraph, has_options):
    children = paragraph.getchildren()
    print(children , '124')
    exit()
###------------------------------------
path = 'src/2019年全国I卷理科数学高考真题.docx'
doc = docx.Document(path)
data_list = processPaper(doc)

choice_que = []
completion_que = []
answer_que = []

index = 0
for dati in data_list:

    i=0
    dati_row=dati[0]
    ti={}
    while(i<len(dati[1])):
        title_indexs=dati[1][i]['title']
        option_indexs=dati[1][i]['options']
        htmls=''
        for index in title_indexs:
            html=paraghraph2html(doc.paragraphs[index],False)
            htmls=htmls+html
        ti['title']=htmls.copy()
        for index in option_indexs:
            html=paraghraph2html(doc.paragraphs[index], True)

            htmls=htmls+html



