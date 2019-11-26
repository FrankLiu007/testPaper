import docx
import re
from lxml import etree
import uuid
from docx_utils.namespaces import namespaces as docx_nsmap
#####主要是进行版面分析，把每个题的标题、选项等部分所在的段落号，计算出来

# 判断是否为标题
def isNumber(char):
    num1 = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9']
    num2 = ['一', '二', '三', '四', '五', '六', '七', '八', '九']
    if char in num1 or char in num2:
        return True
    return False

##获取下一个数字
def get_next_number(text):
    num1 = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9']
    num2 = ['一', '二', '三', '四', '五', '六', '七', '八', '九']
    next_number = ''
    if text.isdigit():
        next_number = str(int(text) + 1)
    else:
        for i in range(0, len(num2)):
            if num2[i] == text:
                next_number = num2[i + 1]
    return next_number

###计算模式字符串
def get_mode_string(text):
    tt = text[0:3]
    k = 0
    mode_string = ''
    for i in range(0, 3):
        if isNumber(text[i]):
            k = i
            break
    if k == 0:
        if text[1] in ['\t', ' ']:
            mode_string = ''
    return mode_string

#分析试卷的版面
# 大题模式为中文数字一 ， 小题模式为 阿拉伯数字 1
def analys_layout(doc):
    '''

    :param doc:
    :return: tree
    格式：[ (row,paragraphs[row].text, mode_text),(...)]
    '''
    pars = []  ###可能
    paragraphs = doc.paragraphs
    row = 0

    ##找出所有带数字的行和行号(  中文数字 一， 阿拉伯数字 1 )
    while (row < len(paragraphs)):
        text = paragraphs[row].text.strip()
        for n in range(0, 4):
            if n >= len(text):
                continue
            if isNumber(text[n]):
                pars.append((n, row, text))  ##n--数字出现的位置，i--段落号
                break
        row = row + 1

    ##找所有模式，行号为 1 或者 一
    modes = []
    for n, row, text in pars:
        if n == 0:
            if text[n] in ['1', '一'] and (not isNumber(text[n + 1])):
                modes.append((n, row, text, text[n]))
        else:
            if text[n] in ['1', '一'] and (not isNumber(text[n - 1])) and (not isNumber(text[n + 1])):
                modes.append((n, row, text, text[n]))

    tree = []
    for i in range(0, len(modes)):  ##每个模式
        n, begin_row, text, mode_text = modes[i]  ##n为行号
        tmp_mode = []
        tmp_mode.append((begin_row, text, mode_text))

        next_number = get_next_number(mode_text)
        if text[n+1]=='．':
            text=text[:n+1]+'.'+text[n+2:] ###全角半角“.”的转换，防止用户没有正确使用“.”号

        for row in range(begin_row + 1, len(paragraphs)):  ###查找某个模式的所有的值
            tt=paragraphs[row].text.strip()
            if len(tt)<n+2:
                continue
            x=tt.find('．')
            if x!=-1 and x <=n+2:
                tt=tt[:x]+'.'+tt[x+1:]   ###全角半角"."的转换，防止用户没有正确使用"."号

            if tt.startswith(text[:n] + mode_text + text[n + 1]):  ##发现了相同的模式，退出
                break
            if tt.startswith(text[:n] + next_number + text[n + 1]):
                # print('next_string:', text[:n] + next_number + text[n + 1], paragraphs[row].text)
                next_number = get_next_number(next_number)
                tmp_mode.append((row, tt, text[:n] + mode_text + text[n + 1]))
        if len(tmp_mode) > 1:
            tree.append(tmp_mode.copy())

    return tree


##获取1个选项，[A-G]. 形式的
def get_option(text):
    text = text.strip()
    indexes = []
    options = []
    for item in re.finditer(r'[A-G][\.．]', text):
        indexes.append((item.group(), item.span()))
    print('in get_option,text=', text)
    if indexes[0][1] != (0, 2):  ###校检结果，保证
        print('获取选择题选项出错，请检查试题格式')
        print('text=', text)

    i = 0
    while (i < len(indexes)):
        b = indexes[i][1][0]
        if i == len(indexes) - 1:
            options.append((option_text[0], text[b:].strip()))
            break
        e = indexes[i + 1][1][0]
        option_text = indexes[i][0]
        options.append((option_text[0], text[b:e].strip()))
        b = e
        i = i + 1
    return options


def parse_one_titype(curr_row, next_row, xiaoti_indexes, paragraphs):  ##处理1个大题，例如 “一、选择题”
    tis = []
    i = 0

    while (i < len(xiaoti_indexes)):
        # r,text,mode_text=xiaoti_indexes[i]
        if xiaoti_indexes[i][0] > curr_row:
            if i == len(xiaoti_indexes) - 1:   ###处理最后一个小题
                ti = parse_ti(xiaoti_indexes, xiaoti_indexes[i][0], next_row, paragraphs)

                tis.append(ti)
                break

            if xiaoti_indexes[i + 1][0] < next_row:
                ti = parse_ti(xiaoti_indexes, xiaoti_indexes[i][0], xiaoti_indexes[i + 1][0], paragraphs)
            else:
                ti = parse_ti(xiaoti_indexes, xiaoti_indexes[i][0], next_row, paragraphs)
                tis.append(ti)
                break
            tis.append(ti)
        i = i + 1

    return tis


####处理1个题
def isObjective(curr_row, next_row, paragraphs):
    # print('next_row=',next_row)
    for i in range(curr_row, next_row):
        text = paragraphs[i].text.strip()
        if re.match(r'[A-G][．\.]', text):
            return (True, i)
    return (False, -1)


####解析1道题
def parse_ti(xiaoti_indexes, curr_row, next_row, paragraphs):
    # curr_row=xiaoti_indexes
    objective, index = isObjective(curr_row, next_row, paragraphs)
    ti = {}
    ti['title'] = []
    if objective:
        options = []
        for i in range(curr_row, index):
            ti['title'].append(i)
        for j in range(index, next_row):
            if re.match(r'[A-G][．\.]', paragraphs[j].text):
                options.append(j)
        ti['options'] = options
    else:
        for i in range(curr_row, next_row):
            ti['title'].append(i)

    return ti

###计算主要模式在tree的位置
##2019.9.30,主模式确定后，副模式应该在主模式之前（行号更小）
def get_main_modes(tree):
    data = []
    i = 0
    while (i < len(tree)):
        data.append((i, len(tree[i]), tree[i][0][2], tree[i][0][0]))  ##i为
        i = i + 1
    data.sort(key=lambda x: x[1], reverse=True)
    print('data=', data)

    primary_mode_index = data[0][0]  ###最长的肯定是主模式，
    primary_mode_text = data[0][2]
    min_row = data[0][-1]
    second_mode_index = data[1][0]

    if primary_mode_text[0] != '1':
        print('试卷格式可能有问题')
        print('模式字符串是：', primary_mode_text)

    i = 1
    while (i < len(data)):
        if data[i][2][0] == primary_mode_text[0]:
            i = i + 1
            continue
        if '一' in data[i][2] and data[i][-1] < min_row:
            second_mode_index = data[i][0]
            break
        i = i + 1

    return (second_mode_index, primary_mode_index)


def check_run(child):
    i = 0
    # print('child=', child)
    run = child.__copy__()
    rPr = run.xpath('.//w:rPr', namespaces=run.nsmap)  ##删除run的属性

    if len(rPr) > 1:
        print('docx格式出错了，len(w:rPr)!=1')
    elif len(rPr) == 1:
        run.remove(rPr[0])
    # else: 没有rPr,啥都不用做

    wt = run.xpath('.//w:t', namespaces=run.nsmap)
    wdrawing = run.xpath('.//w:drawing', namespaces=run.nsmap)
    moMath = run.xpath('.//m:oMath', namespaces=docx_nsmap)
    i = len(wt) + len(wdrawing) + len(moMath)
    return i


###删除空白行
def remove_blank_paragraph(doc):
    is_blank = True

    i = len(doc.paragraphs) - 1
    nn = 0  ##number of blank paragraphs
    while (i >= 0):
        text = ''
        paragraph = doc.paragraphs[i]
        for run in paragraph.runs:
            wdrawing = run.xpath('.//w:drawing', namespaces=run.nsmap)
            if len(wdrawing) > 0:
                is_blank = False
                break
            moMath = run.xpath('.//m:oMath', namespaces=docx_nsmap)
            if len(moMath) > 0:
                is_blank = False
                break

        if paragraph.text.strip() == '':
            is_blank = True
        if is_blank:
            doc.paragraphs.remove(paragraph)

            nn = nn + 1
        i = i - 1
    print('空白段落总数：', nn)

    return 0
##new version of analys_layout
def find_xiaoti_row(dati_start_row,doc):
    xiaoti_row=[]

    paragraphs=doc.paragraphs
    curr_number=0
    for n in range(dati_start_row,len(doc.paragraphs)):
        text=paragraphs[n].text.strip()
        x=''
        if len(text)<3:
            continue
        for i in range(0,4):
            if text[i].isdecimal():
                x=x+text[i]
            else:
                break
        if x and re.match(r'^\d{1,2}[\s\.、．]', text):
            if curr_number!=int(x)-1:
                print(text)
                print('in find_xiaoti_row： 格式识别错误！')
                continue
            curr_number = curr_number + 1
            xiaoti_row.append((n , text,x ))

    return xiaoti_row

####找出大题的行
def   find_dati_row( doc):
    paragraphs=doc.paragraphs
    num2 = {'一':'1','二':'2','三':'3','四':'4','五':'5','六':'6','七':'7','八':'8','九':'9', '十':'10'}

    mode_text=''
    dati_row=[]
    ###get mode text
    for i in range(0, len(paragraphs)):
        text=paragraphs[i].text.strip()
        x=''
        if len(text)<4:
            continue
        for j in range(0,4):
            if text[j] in num2:
                x=x+text[j]
                if mode_text=='' and x=='一':
                    mode_text=text[:j+2]
                    break
        if mode_text:
            # mode.append((i, mode_text))
            break
    if not mode_text:
        print('没有找到大题特征字符：一')
        exit(1)

    curr_number = '一'
    ss = mode_text
    p = mode_text.find('一')
    for i in range(0, len(paragraphs)):
        text=paragraphs[i].text.strip()

        if re.match(r'[\s\.、.]',mode_text[p+1] ) :  ###（一、一 一. ）等大题模式
            rr= '^'+mode_text[:p] + curr_number + r'[\s\.、.]'
            if re.match(rr, text):
                dati_row.append((i, text, ss,))
                next_number = get_next_number(curr_number)
                ss = mode_text.replace('一', next_number)
                curr_number = next_number
        else:  ##不是（一、一 一. ）等大题模式，比如（第一部分 ）
            if text.startswith(ss):
                dati_row.append((i, ss, text))
                next_number = get_next_number(curr_number)
                ss = mode_text.replace('一', next_number)
                curr_number = next_number

    return dati_row


def processPaper2(doc):
    paragraphs=doc.paragraphs

    dati_indexes=find_dati_row( doc)
    xiaoti_indexes=find_xiaoti_row(dati_indexes[0][0], doc)

    ####获取所有大题的  小题
    tis = []
    curr_row, text, mode_text = dati_indexes[0]
    i = 0
    all_ti = []
    while (i < len(dati_indexes)):  ##处理1种题型
        if i < len(dati_indexes) - 1:
            next_row, next_text, mode_text = dati_indexes[i + 1]
            tis = parse_one_titype(curr_row, next_row, xiaoti_indexes, paragraphs)  ##处理1种题型的所有题目
            # print('tis=', tis)
        else:
            tis = parse_one_titype(curr_row, len(paragraphs), xiaoti_indexes, paragraphs)  ##处理最后一个大题，例如“三、综合题”
            # print('tis=', tis)
        all_ti.append((curr_row, tis))
        i = i + 1
        curr_row = next_row

    return all_ti

def processPaper(doc):
    '''
    默认，试卷，题目大题是 一、 这种形式
    :param doc:
    :return:
    '''
    paragraphs = doc.paragraphs
    tree = analys_layout(doc)

    # 获取一份试卷主要的大题和主干小题的在tree里的索引
    dati_mode_index, xiao_mode_index = get_main_modes(tree)  ##试卷的主要2层模式

    ####获取所有大题的  小题
    tis = []

    dati_indexes = tree[dati_mode_index]
    xiaoti_indexes = tree[xiao_mode_index]

    curr_row, text, mode_text = dati_indexes[0]
    i = 0
    all_ti = []
    while (i < len(dati_indexes)):  ##处理1种题型
        if i < len(dati_indexes) - 1:
            next_row, next_text, mode_text = dati_indexes[i + 1]
            tis = parse_one_titype(curr_row, next_row, xiaoti_indexes, paragraphs)  ##处理1种题型的所有题目
            # print('tis=', tis)
        else:
            tis = parse_one_titype(curr_row, len(paragraphs), xiaoti_indexes, paragraphs)  ##处理最后一个大题，例如“三、综合题”
            # print('tis=', tis)
        all_ti.append((curr_row, tis))
        i = i + 1
        curr_row = next_row

    return all_ti


####删除括号及其里面的内容
def remove_brackets(sentence):
    result = re.findall(r'(^.*)[（\(].*[\)）](.*)', sentence)
    return ''.join(result)


'''
试卷的格式，我们认为只有2级，
1.大题（一、填空题）
2. 小题（19.）
'''

if __name__ == "__main__":
    path = 'src/2019年全国II卷文科综合高考真题.docx'
    doc = docx.Document(path)
    all_ti_index = processPaper(doc)
