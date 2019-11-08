import docx
from docx_utils.parse_paper import  processPaper
from docx_utils.ti2html import  get_ti_content
import  re
import json
import sys
import os
'''
###获取答案
#答案和试卷不一样，答案一定是先有题号，然后跟着答案。
# 试卷则不同，试卷的选择题里面有材料题（看一段材料，做几个选择题，请参照文综试卷的选择题部分）
'''
def get_answer(doc , all_indexes):
    all_ans = {}
    for dati in all_indexes:
        curr_dati_row = dati[0]
        xiaoti_indexes = dati[1]
        curr_index = 0
        while curr_index<len(xiaoti_indexes):
            curr_index , ans = get_ti_content(doc , xiaoti_indexes , curr_index ,curr_dati_row, '')
            if len(ans['questions'])!=1:
                print('答案格式错误！')
            content = re.sub(r'\d{1,2}[.．]\s{0,}', '', ans['questions'][0]['stem'])

            all_ans[ans['questions'][0]['index']]=content
    return all_ans

def merge_answer(tis , answer_list):
    for ti in tis:
        reference=''
        for q in ti['questions']:
            q['solution']=answer_list[q['index']]
            reference=reference + q['index']+'. '+ q['solution'] + '</br>'
        ti['reference']=reference
    return 0

#####给题目增加分数和题型
def add_score_and_titype(ti, text):
    score = re.findall(r'每小题(\d{1,2})分', text)     ##，每个小题的分数

    ###题型 titpye
    xx=re.findall(r'(.{1,8}题)', text)[0]
    b=text.find('、')
    e=text.find('题')
    if b!=-1 and e!=-1:
        type_str=xx[b+1:e+1]
        if b > 3:
            print('题型识别可能出错：', xx[b+1:e+1])
    ti['category'] = type_str

    q_tpye=''
    if '只有一项' in text  or '的一项是' in text :
        q_tpye='single'

    if score:
        ti['score'] = int(score[0]) * len(ti['questions'])
    ss=0
    for q in ti['questions']:
        if (not 'type' in q) or  (q_tpye==''):
            q['type']=q_tpye
        if score:
            q['score'] = int(score[0])
        else:
            s = re.findall(r'["(","（"](\d{1,2})分[")","）"]', q['stem'])
            q['score'] = int(s[0])
            ss = ss + q['score']
    if not score:
        ti['score'] = ss
    return 0
def get_tis(all_ti_index):
    mode_text = r'完成\d{1,2}～\d{1,2}题'  ##模式字符串
    tis = []
    # while(i<len(all_ti_index)):
    for dati in all_ti_index:
        curr_dati_row = dati[0]

        xiaoti_indexes = dati[1]          ####题目集合index
        curr_row = xiaoti_indexes[0]      # 题型段落index
        curr_index = 0
        while (curr_index < len(xiaoti_indexes)):

            curr_index, ti = get_ti_content(doc, xiaoti_indexes, curr_index, curr_dati_row, mode_text)
            add_score_and_titype(ti, paragraphs[curr_dati_row].text)
            tis.append(ti)
    return  tis
#-----------------------------------------
def decide_titype(text):
    pass
if __name__ == "__main__":

    if len(sys.argv)==5:
        paper_path=sys.argv[1]
        answer_path=sys.argv[2]

    else:
        print('参数错误，正确用法： process_3x_paper.py 真题.docx 答案.docx')
        paper_path = 'src/2019年全国II卷文科综合高考真题.docx'
        answer_path = 'src/2019年全国II卷文科综合高考真题-答案.docx'


    doc = docx.Document(paper_path)
    all_ti_index = processPaper(doc)
    paragraphs = doc.paragraphs

###处理试卷
    i = 0
    tis=get_tis(all_ti_index)

####处理答案

    doc = docx.Document(answer_path)
    all_answer_index = processPaper(doc)
    answers=get_answer(doc,all_answer_index )
    merge_answer(tis, answers)

    out_path=os.path.splitext(paper_path)[0]+'.json'
    with  open(out_path, 'w', encoding='utf-8') as fp:
        json.dump(tis, fp, ensure_ascii=False,indent = 4, separators=(',', ': '))
