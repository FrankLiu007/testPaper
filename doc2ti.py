import docx

import re

import  json

def get_title(doc , curr_index , e_index):
    paragraphs = doc.paragraphs
    title = ''
    while(1):
        tt = paragraphs[curr_index].text
        tt = tt.strip()

        if not e_index and re.match(r'^([A-G]|[0-9]+)\..*', tt):
            title = title[:-5]
            return title ,curr_index

        if e_index and curr_index>e_index:
            title = title[:-5]
            return title, curr_index

        title += '<span>' + paragraphs[curr_index].text + '</span><br/>'
        curr_index += 1

def get_subject(tt):
    stem = re.findall(r'^[0-9]{1,2}\.(.*)$', tt)
    stem = '<span>' + stem[0].strip() + '</span>'
    return stem

def get_option(tt):
    options = []
    ops = re.findall(r'([A-G]\.[\s]{1}.{1,200})', tt)
    ops = re.split('\s{2,}', ops[0])
    for item in ops:
        label = re.findall(r'([A-G])\.[\s]{1}', item)[0]
        content = re.findall(r'[A-G]\.[\s]{1}(.{1,200})', item)[0]
        temp_option = dict()
        temp_option['label'] = label
        temp_option['content'] = content
        options.append(temp_option)
    return options

def get_temp_que(type , stem , options , one_ans):
    temp_que = dict()
    temp_que['type'] = type
    temp_que['stem'] = stem
    temp_que['options'] = options
    temp_que['solution'] = one_ans
    return temp_que

def get_one_item(title,category, reference, questions):
    ti = dict()
    ti['title'] = title
    ti['category'] = category
    ti['reference'] = '<span>'+ reference +'</span>'
    ti['questions'] = questions
    return ti

def getReading(doc , b_index , e_index , ans_index):
    ans = processAnswer('21', '41' ,'1')
    paragraphs = doc.paragraphs
    curr_paragraph_index = b_index
    type = 'SINGLE'
    category = '阅读理解'
    questions = []
    all_ans = ''
    index = ans_index

    title = get_title(doc , curr_paragraph_index , 0)[0]
    curr_paragraph_index = get_title(doc , curr_paragraph_index , 0)[1]

    options = []
    have_stem = 0

    while(1):
        tt = paragraphs[curr_paragraph_index].text
        tt = tt.strip()
        curr_paragraph_index += 1

        if re.match(r'^[0-9]{1,2}\..*$', tt):     ###找到1个小题
            have_stem = 1
            if not options:
                stem = get_subject(tt)
                continue

            one_ans = ans[index]
            all_ans += one_ans
            temp_que = get_temp_que(type , stem , options , one_ans)
            index += 1
            questions.append(temp_que)

            options=[]
            stem=get_subject(tt)

        if re.match(r'^[A-G]\..*', tt): ###找到一个选项
            temp_option = get_option(tt)
            for item in temp_option:
                options.append(item)

        if curr_paragraph_index>e_index:
            break

    if not have_stem:
        stem = title
        title = ''
        type = 'GENERAL'

        for i in range(0 , 5):
            one_ans = ans[index]
            all_ans+=one_ans
            index += 1
        temp_que = get_temp_que(type , stem , options , '')
        questions.append(temp_que)

    else:
        one_ans = ans[index]
        all_ans += one_ans
        temp_que = get_temp_que(type, stem, options, one_ans)
        questions.append(temp_que)

    ti = get_one_item(title ,category, all_ans , questions)

    return ti

def getApplication(doc , b_index , e_index , ans_index):
    ans = processAnswer('41', '61' , '1')
    all_ans = ''

    paragraphs = doc.paragraphs
    curr_paragraph_index = b_index
    questions = []
    category = '语言知识运用'

    title = get_title(doc , curr_paragraph_index , 0)[0]
    curr_paragraph_index = get_title(doc , curr_paragraph_index , 0)[1]

    while(1):
        tt = paragraphs[curr_paragraph_index].text
        tt = tt.strip()
        curr_paragraph_index += 1

        if e_index < curr_paragraph_index-1:
            break

        # 语言综合运用部分，每个小题的title为''
        if re.match(r'^[0-9]{1,2}\..*$', tt):     ###找到1个小题
            stem = ''
            type = 'SINGLE'
            options = get_option(tt)
            one_ans = ans[ans_index]
            all_ans+=ans[ans_index]
            temp_que = get_temp_que(type, stem, options, one_ans)
            ans_index += 1
            questions.append(temp_que)

    ti = get_one_item(title,category, all_ans, questions)

    return ti

def getWriting(doc , b_index , e_index , ans):
    category = doc.paragraphs[b_index].text
    b_index += 1
    index = category.find('（')
    category = category[5: index]  # 题目类型
    type = 'GENERAL'
    title = ''
    options = []
    solution = ''
    questions = []

    stem = get_title(doc , b_index , e_index)[0]

    temp_que = get_temp_que(type , stem , options , solution)

    questions.append(temp_que)

    ti = get_one_item(title,category, ans, questions)

    return ti

##### 处理阅读理解------
def processReadings(doc,ti_index, curr_paragraph_index):
    b_index=0
    e_index=0
    ans_index = 0
    all_readings=[]
    curr_paragraph_index += 1
    while (1):
        curr_paragraph_index += 1
        tt1 = doc.paragraphs[curr_paragraph_index].text
        tt1=tt1.strip()

        if "部分" in tt1 and "第" in tt1:
            b_index += 1
            e_index = curr_paragraph_index -1
            stem = get_title(doc , b_index , e_index)[0]
            ans = processAnswer('36', '41', '1')
            temp_que = get_temp_que('GENERAL' , stem , [] , '')
            question = []
            question.append(temp_que)
            result = get_one_item('' , '阅读理解' , ans ,question )
            all_readings.append(result)
            break

        if (tt1 in ['A','B','C','D']) or ('第二节' in tt1):
            if tt1 == 'A':
                b_index = curr_paragraph_index + 1
                continue
            e_index = curr_paragraph_index -1
            result = getReading(doc, b_index, e_index , ans_index)
            ans_index += len(result['questions'])
            all_readings.append(result)
            b_index = curr_paragraph_index+1
            continue

    return all_readings

##### 处理语言知识运用------
def processApplication(doc,ti_index, curr_paragraph_index):
    b_index = 0
    ans_index = 0
    all_application = []
    category = '语言知识运用'
    while(1):
        curr_paragraph_index += 1
        tt = doc.paragraphs[curr_paragraph_index].text
        if '第四部分' in tt :
            e_index = curr_paragraph_index - 1
            title = get_title(doc, b_index, e_index)[0]
            ans = processAnswer('61', '71' , '2')
            reference = ''
            for item in ans:
                item = '  '+item+'  '
                reference += item

            type = 'GENERAL'
            stem = ''
            options = []
            solution = ''
            temp_que = get_temp_que(type, stem, options, solution)
            questions = []
            questions.append(temp_que)

            ti = get_one_item(title,category, reference, questions)
            all_application.append(ti)

            break

        if '第二节' in tt:
            e_index = curr_paragraph_index-1
            result = getApplication(doc , b_index , e_index , ans_index)
            all_application.append(result)
            b_index = curr_paragraph_index + 2

        if '第一节' in tt:
            b_index = curr_paragraph_index+2

    return all_application

##### 处理写作------
def processWriting(doc,ti_index, curr_paragraph_index):
    b_index = 0
    all_writing = []

    while(1):
        curr_paragraph_index += 1
        tt = doc.paragraphs[curr_paragraph_index].text

        if (len(tt) == 0) or '第二节' in tt:
            if (len(tt) == 0):
                ans = processAnswer('72', 0, '3')
            if '第二节' in tt:
                ans = processAnswer('71', '72', '3')
            e_index = curr_paragraph_index-1
            result =  getWriting(doc , b_index , e_index , ans)
            all_writing.append(result)
            b_index = curr_paragraph_index
            if (len(tt) == 0):
                break
            continue

        if '第一节' in tt:
            b_index = curr_paragraph_index

    return all_writing

##### 处理答案------
def processAnswer(b_ans , e_ans , type):
    b_index = answer.index(b_ans)
    if not e_ans:
      ans = answer[b_index: ]
    else:
        e_index = answer.index(e_ans)
        ans = answer[b_index: e_index]

    if type in '1':
        ans = re.findall(r"[0-9-]+\s*([A-Z]+)", ans)
        ans = ''.join(ans)
    elif type in '2':
        ans_list= re.split(r"\s{2,}" , ans)
        ans = []
        for item in ans_list:
            if item:
                ans.append(re.findall(r'[0-9]+\.(.+)' , item)[0])
    else:
        ans = re.findall(r"[0-9]+\.(.+)" , ans)[0]

    return  ans


path="src/（精校版）2018年全国卷Ⅰ英语高考真题文档版（含答案）.docx"
doc=docx.Document(path)
x=1
# 获得全部的段落
print("len(paragraphs)=", len(doc.paragraphs))
curr_ti_index=1
ti=[]
curr_paragraph_index=-1
ti_index=0

answer = ''
all_data = []
all_reading = []
all_application = []
all_writing = []

now_paragraph_index = curr_paragraph_index +1
while(1):
    now_paragraph_index += 1
    tt = doc.paragraphs[now_paragraph_index].text
    tt = tt.strip()

    if '试卷答案' in tt:
        b_index = now_paragraph_index + 1
        e_index = 0
        while (1):
            now_paragraph_index += 1
            tt = doc.paragraphs[now_paragraph_index].text
            answer += tt
            if now_paragraph_index+1 == len(doc.paragraphs):
                break
        answer  = answer.replace('\t' , '')
        break

while(1):
    curr_paragraph_index += 1
    tt=doc.paragraphs[curr_paragraph_index].text
    tt=tt.strip()
    if tt.startswith("第一部分")  and "听力" in tt:
        # curr_paragraph_index=processListenings(doc,ti_index, curr_paragraph_index)
        continue
    elif tt.startswith("第二部分")  and "阅读理解" in tt:
        all_reading=processReadings(doc,ti_index, curr_paragraph_index)
        continue
    elif tt.startswith("第三部分")  and "语言知识运用" in tt:
        all_application = processApplication(doc , ti_index , curr_paragraph_index)
        continue
    elif tt.startswith('第四部分') and '写作' in tt:
        all_writing = processWriting(doc , ti_index , curr_paragraph_index)
        break

all_data.extend(all_reading)
all_data.extend(all_application)
all_data.extend(all_writing)

fp = open('data.json' , 'w' , encoding='utf-8')
json.dump(all_data , fp)
fp.close()








